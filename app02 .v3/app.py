import os
from flask import Flask, request, send_file, send_from_directory
from docx import Document
import io

app = Flask(__name__)

def preencher_template(dados, template_filename):
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", template_filename)
    print(f"Usando o template em: {template_path}")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"O arquivo {template_path} não foi encontrado.")
    
    doc = Document(template_path)
    
    # Substituir os placeholders no documento
    for paragrafo in doc.paragraphs:
        for placeholder, valor in dados.items():
            if isinstance(valor, list):
                valor = ', '.join(valor)  # Convertendo listas para strings
            if placeholder in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(placeholder, valor)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for placeholder, valor in dados.items():
                        if isinstance(valor, list):
                            valor = ', '.join(valor)  # Convertendo listas para strings
                        if placeholder in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(placeholder, valor)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/')
def formulario():
    return send_from_directory('static', 'index.html')

@app.route('/gerar', methods=['POST'])
def gerar_documento():
    aws_account_id = request.form['aws_account_id']
    nome_cliente = request.form['nome_cliente']
    plano = request.form['plano']
    software = request.form['software']
    inicio_em_mes = request.form['inicio_em_mes']
    inicio_em_ano = request.form['inicio_em_ano']
    link_grupo_whatsapp = request.form['link_grupo_whatsapp']
    link_grupo_telegram = request.form.get('link_grupo_telegram', '')  # Usando request.form.get() para evitar KeyError
    link_ims = request.form['link_ims']
    qtde_servidores = request.form['qtde_servidores']
    rto_parcial = request.form['rto_parcial']
    rto_completo = request.form['rto_completo']
    servicos_criticos = request.form.get('servicos_criticos', '')
    
    inicio_em = f"{inicio_em_mes}/{inicio_em_ano}"

    # Selecionar o template correto com base no plano
    if plano == "Vault":
        template_filename = "sunny_vault.docx"
    elif plano == "Start":
        template_filename = "sunny_start.docx"
    elif plano == "Essentials":
        template_filename = "sunny_essentials.docx"
    elif plano == "Enterprise":
        template_filename = "sunny_enterprise.docx"
    else:
        return "Plano inválido", 400

    dados_cliente = {
        "{aws_account_id}": aws_account_id,
        "{nome_cliente}": nome_cliente,
        "{plano}": plano,
        "{software}": software,
        "{inicio_em}": inicio_em,
        "{link_grupo_whatsapp}": link_grupo_whatsapp,
        "{link_grupo_telegram}": link_grupo_telegram,
        "{link_ims}": link_ims,
        "{qtde_servidores}": qtde_servidores,
        "{rto_parcial}": rto_parcial,
        "{rto_completo}": rto_completo,
        "{servicos_criticos}": servicos_criticos,
    }
    
    output = preencher_template(dados_cliente, template_filename)
    
    download_name = f'Sunny {plano} - {nome_cliente}.docx'
    
    return send_file(output, as_attachment=True, download_name=download_name)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
